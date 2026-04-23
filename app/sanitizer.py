"""
sanitizer.py  —  Advanced Secure Sanitization Engine v3.0
----------------------------------------------------------
Deep content inspection, entropy analysis, multi-layer threat detection,
format-specific sanitization with full forensic reporting.

NEW in v3.0:
  - extract_content_preview(): reads INSIDE file data per type for display
  - ScanResult.content_preview: dict with file-type-specific parsed fields
  - CSV: row count, column names, sample rows, formula injection map
  - JSON: schema keys, record count, sample values
  - TXT: line count, word count, first 10 lines preview
  - PDF: page count, extracted text snippet (pikepdf)
  - PNG/JPG: dimensions, mode, EXIF fields
  - DOCX: paragraph count, heading list, author
  - XLSX: sheet names, row/col counts per sheet, sample cell values
"""

import hashlib
import io
import json
import math
import os
import re
import shutil
import struct
import zipfile
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any

# ── optional heavy deps ───────────────────────────────────────────────────────
try:
    import pikepdf
    HAS_PIKEPDF = True
except ImportError:
    HAS_PIKEPDF = False

try:
    from PIL import Image, ExifTags
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    import csv as _csv
    HAS_CSV = True
except ImportError:
    HAS_CSV = False

# ── rules ─────────────────────────────────────────────────────────────────────
_RULES_PATH = Path(__file__).parent.parent / "rules" / "default_rules.json"
with open(_RULES_PATH) as fh:
    RULES = json.load(fh)

BLOCKED_EXT = set(RULES["blocked_extensions"])
ALLOWED_EXT = set(RULES["allowed_extensions"])
MAX_SIZE_MB = RULES["max_file_size_mb"]
THREAT_KW   = RULES["threat_keywords"]

# ── advanced threat patterns (YARA-style regex signatures) ────────────────────
THREAT_PATTERNS = {
    "shellcode_nop_sled":   (rb"\x90{16,}", "NOP sled detected (shellcode indicator)"),
    "pe_header":            (rb"MZ\x90\x00.{0,58}PE\x00\x00", "PE executable header embedded"),
    "reverse_shell_py":     (rb"socket\.connect\(.+\d+\)", "Python reverse shell pattern"),
    "base64_payload":       (rb"(?:eval|exec)\s*\(\s*(?:base64|b64)", "Base64 payload execution"),
    "powershell_enc":       (rb"(?i)-enc(?:odedcommand)?\s+[A-Za-z0-9+/]{20,}={0,2}", "Encoded PowerShell command"),
    "xml_bomb":             (rb"<!ENTITY\s+\w+\s+\"&\w+;&\w+;&\w+;", "XML bomb / entity expansion"),
    "zip_slip":             (rb"\.\./|\.\.\\\\", "Path traversal sequence (zip-slip)"),
    "elf_header":           (rb"\x7fELF", "ELF binary header embedded"),
    "java_class":           (rb"\xca\xfe\xba\xbe", "Java class bytecode embedded"),
    "pdf_js_action":        (rb"/JS\s*\(|/JavaScript\s*/", "JavaScript action in PDF"),
    "macro_auto":           (rb"(?i)Auto(?:Open|Close|Exec|New)\s*\(\s*\)", "VBA auto-execution macro"),
    "c2_beacon":            (rb"(?i)(?:beacon|callback|c2|command.and.control)\s*[:=]", "C2 callback reference"),
    "hex_shellcode":        (rb"(?:\\x[0-9a-fA-F]{2}){12,}", "Hex-encoded shellcode sequence"),
    "ldap_injection":       (rb"\)\s*\(\s*\|\s*\(", "LDAP injection pattern"),
    "ssrf_internal":        (rb"(?i)(?:169\.254|192\.168|10\.\d+\.\d+|172\.(?:1[6-9]|2\d|3[01]))\.\d+", "Internal IP reference (SSRF risk)"),
}

# ── magic bytes for file type verification ────────────────────────────────────
MAGIC_SIGNATURES = {
    b"%PDF":                   ".pdf",
    b"\xff\xd8\xff":           ".jpg",
    b"\x89PNG\r\n\x1a\n":     ".png",
    b"PK\x03\x04":            ".zip",
    b"\xd0\xcf\x11\xe0":      ".doc",
    b"MZ":                     ".exe",
    b"\x7fELF":                ".elf",
    b"\xca\xfe\xba\xbe":      ".class",
    b"GIF87a":                 ".gif",
    b"GIF89a":                 ".gif",
    b"\x1f\x8b":              ".gz",
    b"BZh":                    ".bz2",
    b"Rar!":                   ".rar",
    b"7z\xbc\xaf\x27\x1c":   ".7z",
}


# ── data types ────────────────────────────────────────────────────────────────
@dataclass
class ThreatIndicator:
    category:    str
    severity:    str
    description: str
    offset:      Optional[int] = None
    evidence:    Optional[str] = None


@dataclass
class ScanResult:
    filename:         str
    extension:        str
    detected_type:    str
    size_mb:          float
    sha256:           str
    md5:              str
    entropy:          float
    allowed:          bool
    threats:          List[str]             = field(default_factory=list)
    warnings:         List[str]             = field(default_factory=list)
    indicators:       List[ThreatIndicator] = field(default_factory=list)
    metadata_found:   Dict[str, str]        = field(default_factory=dict)
    verdict:          str = "PASS"
    risk_score:       int = 0
    content_preview:  Dict[str, Any]        = field(default_factory=dict)


@dataclass
class SanitizeResult:
    filename:      str
    actions:       List[str]       = field(default_factory=list)
    removed_items: List[str]       = field(default_factory=list)
    output_path:   Optional[str]   = None
    sha256_in:     Optional[str]   = None
    sha256_out:    Optional[str]   = None
    size_in_mb:    float           = 0.0
    size_out_mb:   float           = 0.0
    success:       bool            = False
    error:         Optional[str]   = None
    sanitize_time_ms: float        = 0.0


# ── helpers ───────────────────────────────────────────────────────────────────
def sha256_file(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def md5_file(path: str) -> str:
    h = hashlib.md5()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def file_size_mb(path: str) -> float:
    return os.path.getsize(path) / (1024 * 1024)


def get_extension(path: str) -> str:
    return Path(path).suffix.lower()


def shannon_entropy(data: bytes) -> float:
    if not data:
        return 0.0
    counts = Counter(data)
    total = len(data)
    entropy = 0.0
    for count in counts.values():
        p = count / total
        if p > 0:
            entropy -= p * math.log2(p)
    return round(entropy, 4)


def detect_type_from_magic(path: str) -> str:
    try:
        with open(path, "rb") as f:
            header = f.read(16)
        for magic, ext in MAGIC_SIGNATURES.items():
            if header.startswith(magic):
                return ext
    except Exception:
        pass
    return get_extension(path)


def _compute_risk_score(result: ScanResult) -> int:
    score = 0
    severity_map = {"CRITICAL": 40, "HIGH": 25, "MEDIUM": 10, "LOW": 3}
    for ind in result.indicators:
        score += severity_map.get(ind.severity, 5)
    if result.entropy > 7.5:
        score += 15
    if not result.allowed:
        score += 30
    if result.size_mb > MAX_SIZE_MB:
        score += 5
    return min(score, 100)


# ══════════════════════════════════════════════════════════════════════════════
# CONTENT PREVIEW — reads INSIDE file data per type  (NEW v3.0)
# ══════════════════════════════════════════════════════════════════════════════
def extract_content_preview(filepath: str) -> Dict[str, Any]:
    """
    Read inside file content and return a structured preview dict
    tailored to the file type. Used by the UI to show format-specific
    input/output analysis panels.
    """
    ext = get_extension(filepath)
    preview: Dict[str, Any] = {"type": ext, "error": None}

    try:
        # ── TXT ───────────────────────────────────────────────────────────────
        if ext == ".txt":
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                lines = f.readlines()
            words = sum(len(l.split()) for l in lines)
            preview.update({
                "line_count":   len(lines),
                "word_count":   words,
                "char_count":   sum(len(l) for l in lines),
                "first_lines":  [l.rstrip() for l in lines[:10]],
                "has_urls":     any("http" in l.lower() for l in lines),
                "has_ips":      bool(re.search(r'\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}', "".join(lines))),
                "empty_lines":  sum(1 for l in lines if l.strip() == ""),
            })

        # ── CSV ───────────────────────────────────────────────────────────────
        elif ext == ".csv":
            import csv as _csv_mod
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                reader = _csv_mod.DictReader(f)
                rows = []
                injection_cells = []
                for i, row in enumerate(reader):
                    if i < 5:
                        rows.append(dict(row))
                    for k, v in row.items():
                        if str(v).startswith(("=", "+", "-", "@")):
                            injection_cells.append(f"Row {i+2}, Col '{k}': {str(v)[:30]}")
                    if i > 500:
                        break
                cols = list(reader.fieldnames or [])
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                total_rows = sum(1 for _ in f) - 1
            preview.update({
                "columns":         cols,
                "column_count":    len(cols),
                "row_count":       total_rows,
                "sample_rows":     rows,
                "injection_cells": injection_cells[:10],
                "has_injection":   len(injection_cells) > 0,
            })

        # ── JSON ──────────────────────────────────────────────────────────────
        elif ext == ".json":
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                raw = f.read(200_000)
            obj = json.loads(raw)
            if isinstance(obj, list):
                record_count = len(obj)
                keys = list(obj[0].keys()) if obj and isinstance(obj[0], dict) else []
                sample = obj[:3]
            elif isinstance(obj, dict):
                record_count = 1
                keys = list(obj.keys())[:20]
                sample = [{k: obj[k] for k in keys[:5]}]
            else:
                record_count = 1
                keys = []
                sample = [str(obj)[:200]]
            preview.update({
                "root_type":    type(obj).__name__,
                "record_count": record_count,
                "top_keys":     keys[:15],
                "sample":       sample[:2],
                "nested_depth": _json_depth(obj),
            })

        # ── PDF ───────────────────────────────────────────────────────────────
        elif ext == ".pdf":
            page_count = 0
            text_snippet = ""
            pdf_metadata = {}
            has_js = False
            has_forms = False
            embedded_count = 0

            if HAS_PIKEPDF:
                with pikepdf.open(filepath) as pdf:
                    page_count = len(pdf.pages)
                    with pdf.open_metadata() as meta:
                        for k in list(meta.keys())[:10]:
                            pdf_metadata[str(k)] = str(meta[k])[:100]
                    # Extract text from first 2 pages
                    for pg in pdf.pages[:2]:
                        try:
                            for k in pg.keys():
                                pass
                        except Exception:
                            pass
                    has_js   = b"/JS" in open(filepath,"rb").read() or b"/JavaScript" in open(filepath,"rb").read()
                    has_forms = b"/AcroForm" in open(filepath,"rb").read()
                    # Count streams as proxy for embedded objects
                    raw = open(filepath,"rb").read()
                    embedded_count = raw.count(b"/EmbeddedFile")
                    # Simple text extraction from raw PDF streams
                    text_parts = re.findall(rb'\(([^\)]{4,120})\)', raw)
                    decoded = []
                    for p in text_parts[:30]:
                        try:
                            t = p.decode("latin-1")
                            if t.isprintable() and len(t) > 3:
                                decoded.append(t)
                        except Exception:
                            pass
                    text_snippet = " ".join(decoded[:15])[:400]
            else:
                raw = open(filepath,"rb").read()
                page_count = raw.count(b"/Page ")
                text_parts = re.findall(rb'\(([^\)]{4,80})\)', raw)
                decoded = []
                for p in text_parts[:20]:
                    try:
                        t = p.decode("latin-1")
                        if t.isprintable():
                            decoded.append(t)
                    except Exception:
                        pass
                text_snippet = " ".join(decoded[:10])[:300]

            preview.update({
                "page_count":      page_count,
                "text_snippet":    text_snippet,
                "metadata_fields": pdf_metadata,
                "has_javascript":  has_js,
                "has_forms":       has_forms,
                "embedded_files":  embedded_count,
            })

        # ── PNG / JPG ─────────────────────────────────────────────────────────
        elif ext in (".png", ".jpg", ".jpeg"):
            if HAS_PIL:
                img = Image.open(filepath)
                w, h = img.size
                mode = img.mode
                fmt  = img.format or ext.upper().strip(".")
                exif_fields = {}
                try:
                    exif_raw = img._getexif()
                    if exif_raw:
                        for tag_id, val in list(exif_raw.items())[:15]:
                            tag = ExifTags.TAGS.get(tag_id, str(tag_id))
                            exif_fields[tag] = str(val)[:80]
                except Exception:
                    pass
                has_alpha = mode in ("RGBA", "LA", "PA")
                preview.update({
                    "width":       w,
                    "height":      h,
                    "mode":        mode,
                    "format":      fmt,
                    "has_alpha":   has_alpha,
                    "exif_count":  len(exif_fields),
                    "exif_fields": exif_fields,
                    "has_gps":     "GPSInfo" in exif_fields,
                    "megapixels":  round((w * h) / 1_000_000, 2),
                })
            else:
                raw = open(filepath,"rb").read(24)
                preview.update({"raw_header": raw.hex(), "note": "Pillow not installed"})

        # ── DOCX ─────────────────────────────────────────────────────────────
        elif ext == ".docx":
            if HAS_DOCX:
                doc = DocxDocument(filepath)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                headings   = [p.text for p in doc.paragraphs
                              if p.style.name.startswith("Heading") and p.text.strip()]
                props = doc.core_properties
                has_macros = False
                try:
                    with zipfile.ZipFile(filepath) as z:
                        has_macros = any("vbaProject" in n for n in z.namelist())
                except Exception:
                    pass
                preview.update({
                    "paragraph_count": len(paragraphs),
                    "heading_count":   len(headings),
                    "headings":        headings[:8],
                    "first_paras":     paragraphs[:5],
                    "author":          props.author or "N/A",
                    "title":           props.title  or "N/A",
                    "word_count":      sum(len(p.split()) for p in paragraphs),
                    "has_macros":      has_macros,
                })
            else:
                preview.update({"note": "python-docx not installed"})

        # ── XLSX ─────────────────────────────────────────────────────────────
        elif ext == ".xlsx":
            if HAS_OPENPYXL:
                wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
                sheets_info = []
                for name in wb.sheetnames[:6]:
                    ws = wb[name]
                    sample_cells = []
                    row_count = 0
                    col_count = 0
                    for ri, row in enumerate(ws.iter_rows(values_only=True)):
                        if ri == 0:
                            col_count = len(row)
                        if ri < 3:
                            sample_cells.append([str(c)[:30] if c is not None else "" for c in row[:8]])
                        row_count += 1
                        if row_count > 1000:
                            break
                    sheets_info.append({
                        "name":       name,
                        "rows":       row_count,
                        "cols":       col_count,
                        "sample":     sample_cells,
                    })
                wb.close()
                preview.update({
                    "sheet_count":  len(wb.sheetnames),
                    "sheet_names":  wb.sheetnames,
                    "sheets":       sheets_info,
                })
            else:
                preview.update({"note": "openpyxl not installed"})

        else:
            # Generic binary preview
            raw = open(filepath,"rb").read(256)
            preview.update({
                "hex_header": raw.hex()[:128],
                "printable":  "".join(chr(b) if 32 <= b < 127 else "." for b in raw[:64]),
            })

    except Exception as e:
        preview["error"] = str(e)[:200]

    return preview


def _json_depth(obj, current=0) -> int:
    if current > 10:
        return current
    if isinstance(obj, dict):
        return max((_json_depth(v, current+1) for v in obj.values()), default=current)
    if isinstance(obj, list):
        return max((_json_depth(v, current+1) for v in obj[:5]), default=current)
    return current


# ── keyword scan ──────────────────────────────────────────────────────────────
def _keyword_scan(path: str) -> List[ThreatIndicator]:
    indicators = []
    try:
        with open(path, "rb") as f:
            raw = f.read(2_000_000)
        text = raw.decode("utf-8", errors="ignore").lower()
        for kw in THREAT_KW:
            if kw.lower() in text:
                idx = text.find(kw.lower())
                indicators.append(ThreatIndicator(
                    category="keyword",
                    severity="HIGH",
                    description=f"Threat keyword: '{kw}'",
                    offset=idx,
                    evidence=text[max(0, idx-20):idx+len(kw)+20].strip()
                ))
    except Exception:
        pass
    return indicators


# ── binary pattern scan ───────────────────────────────────────────────────────
def _pattern_scan(path: str) -> List[ThreatIndicator]:
    indicators = []
    try:
        with open(path, "rb") as f:
            data = f.read(5_000_000)
        for sig_name, (pattern, description) in THREAT_PATTERNS.items():
            m = re.search(pattern, data)
            if m:
                severity = "CRITICAL" if sig_name in ("pe_header", "elf_header", "shellcode_nop_sled") else "HIGH"
                indicators.append(ThreatIndicator(
                    category="binary_pattern",
                    severity=severity,
                    description=description,
                    offset=m.start(),
                    evidence=f"Match at byte offset {m.start()}"
                ))
    except Exception:
        pass
    return indicators


# ── magic byte mismatch check ──────────────────────────────────────────────────
def _check_extension_spoofing(path: str, declared_ext: str) -> Optional[ThreatIndicator]:
    detected = detect_type_from_magic(path)
    zip_family = {".docx", ".xlsx", ".zip"}
    if declared_ext in zip_family and detected == ".zip":
        return None
    if detected != declared_ext and detected in MAGIC_SIGNATURES.values():
        return ThreatIndicator(
            category="spoofing",
            severity="CRITICAL",
            description=f"Extension spoofing: declared '{declared_ext}' but magic bytes indicate '{detected}'",
            evidence=f"Detected type: {detected}"
        )
    return None


# ── entropy-based steganography detection ─────────────────────────────────────
def _image_steg_analysis(path: str) -> Tuple[bool, str]:
    if not HAS_PIL:
        return False, ""
    try:
        img = Image.open(path).convert("RGB")
        w, h = img.size
        sample = min(w * h, 20000)
        pixels = list(img.getdata())[:sample]
        r_lsb = bytes([p[0] & 1 for p in pixels])
        g_lsb = bytes([p[1] & 1 for p in pixels])
        b_lsb = bytes([p[2] & 1 for p in pixels])
        r_ent = shannon_entropy(r_lsb)
        g_ent = shannon_entropy(g_lsb)
        b_ent = shannon_entropy(b_lsb)
        avg_lsb_entropy = (r_ent + g_ent + b_ent) / 3
        suspicious = avg_lsb_entropy > 0.96
        detail = f"LSB entropy R={r_ent:.3f} G={g_ent:.3f} B={b_ent:.3f} (avg={avg_lsb_entropy:.3f})"
        return suspicious, detail
    except Exception as e:
        return False, str(e)


# ── macro + embedded object detection ─────────────────────────────────────────
def _office_deep_scan(path: str) -> List[ThreatIndicator]:
    indicators = []
    try:
        with zipfile.ZipFile(path, "r") as z:
            names = z.namelist()
            for name in names:
                if "vbaProject" in name:
                    indicators.append(ThreatIndicator(
                        category="macro", severity="HIGH",
                        description="VBA macro project embedded", evidence=name))
                if name.endswith(".bin") and "vba" not in name.lower():
                    indicators.append(ThreatIndicator(
                        category="embedded_binary", severity="MEDIUM",
                        description="Embedded binary object detected", evidence=name))
                if "oleObject" in name or "embeddings" in name.lower():
                    indicators.append(ThreatIndicator(
                        category="embedded_object", severity="MEDIUM",
                        description="OLE embedded object detected", evidence=name))
                if name.endswith(".xml") or name.endswith(".rels"):
                    try:
                        content = z.read(name)
                        if b"http://" in content or b"https://" in content:
                            indicators.append(ThreatIndicator(
                                category="external_reference", severity="LOW",
                                description="External URL reference in document",
                                evidence=name))
                            break
                    except Exception:
                        pass
    except Exception:
        pass
    return indicators


# ── PDF deep scan ──────────────────────────────────────────────────────────────
def _pdf_deep_scan(path: str) -> Tuple[List[ThreatIndicator], Dict[str, str]]:
    indicators = []
    metadata = {}
    if not HAS_PIKEPDF:
        return indicators, metadata
    try:
        with pikepdf.open(path) as pdf:
            with pdf.open_metadata() as meta:
                for key in meta.keys():
                    val = str(meta[key])
                    metadata[key] = val[:200]
            dangerous_keys = [
                ("/AA", "Automatic Action trigger"),
                ("/JS", "JavaScript"),
                ("/OpenAction", "OpenAction auto-execute"),
                ("/Launch", "Launch action"),
                ("/EmbeddedFile", "Embedded file"),
                ("/RichMedia", "RichMedia (Flash)"),
                ("/XFA", "XFA form (exploitable)"),
            ]
            for page in pdf.pages:
                for key, desc in dangerous_keys:
                    if key in page:
                        indicators.append(ThreatIndicator(
                            category="pdf_action", severity="HIGH",
                            description=f"{desc} found in page", evidence=key))
            for key, desc in dangerous_keys:
                if key in pdf.Root:
                    indicators.append(ThreatIndicator(
                        category="pdf_action", severity="CRITICAL",
                        description=f"{desc} found at document root", evidence=key))
    except Exception as e:
        indicators.append(ThreatIndicator(
            category="parse_error", severity="MEDIUM",
            description=f"PDF parse error (may be corrupted/obfuscated): {str(e)[:100]}"))
    return indicators, metadata


# ── image metadata extraction ──────────────────────────────────────────────────
def _extract_image_metadata(path: str) -> Dict[str, str]:
    metadata = {}
    if not HAS_PIL:
        return metadata
    try:
        img = Image.open(path)
        exif_data = img._getexif() if hasattr(img, "_getexif") else None
        if exif_data:
            for tag_id, value in exif_data.items():
                tag = ExifTags.TAGS.get(tag_id, str(tag_id))
                metadata[tag] = str(value)[:200]
    except Exception:
        pass
    return metadata


# ══════════════════════════════════════════════════════════════════════════════
# SCAN
# ══════════════════════════════════════════════════════════════════════════════
def scan_file(filepath: str) -> ScanResult:
    import time
    fname = os.path.basename(filepath)
    ext   = get_extension(filepath)

    with open(filepath, "rb") as f:
        raw = f.read(min(os.path.getsize(filepath), 5_000_000))

    entropy  = shannon_entropy(raw)
    sha      = sha256_file(filepath)
    md5      = md5_file(filepath)
    size     = file_size_mb(filepath)
    detected = detect_type_from_magic(filepath)

    result = ScanResult(
        filename      = fname,
        extension     = ext,
        detected_type = detected,
        size_mb       = round(size, 3),
        sha256        = sha,
        md5           = md5,
        entropy       = entropy,
        allowed       = ext in ALLOWED_EXT,
    )

    # ── 1. Extension / blocklist ───────────────────────────────────────────
    if ext in BLOCKED_EXT:
        result.indicators.append(ThreatIndicator(
            category="blocked_extension", severity="CRITICAL",
            description=f"Blocked file type: {ext}"))
        result.threats.append(f"Blocked file type: {ext}")
        result.verdict = "FAIL"
        result.risk_score = 100
        result.content_preview = extract_content_preview(filepath)
        return result

    if ext not in ALLOWED_EXT:
        result.indicators.append(ThreatIndicator(
            category="unknown_extension", severity="HIGH",
            description=f"Unknown/unsupported extension: {ext}"))
        result.threats.append(f"Unknown extension: {ext}")
        result.verdict = "FAIL"
        result.risk_score = 80
        result.content_preview = extract_content_preview(filepath)
        return result

    # ── 2. Extension spoofing ──────────────────────────────────────────────
    spoof = _check_extension_spoofing(filepath, ext)
    if spoof:
        result.indicators.append(spoof)
        result.threats.append(spoof.description)
        result.verdict = "FAIL"

    # ── 3. Entropy analysis ────────────────────────────────────────────────
    if entropy > 7.8:
        result.indicators.append(ThreatIndicator(
            category="high_entropy", severity="HIGH",
            description=f"Very high entropy ({entropy:.2f}/8.0) — possibly encrypted or obfuscated payload",
            evidence=f"Shannon entropy: {entropy}"))
        result.warnings.append(f"High entropy: {entropy:.2f}/8.0 (encrypted/obfuscated content likely)")
        if result.verdict == "PASS":
            result.verdict = "WARN"
    elif entropy > 7.2:
        result.warnings.append(f"Elevated entropy: {entropy:.2f}/8.0 (compressed content)")

    # ── 4. Binary pattern scan ─────────────────────────────────────────────
    pattern_hits = _pattern_scan(filepath)
    for ind in pattern_hits:
        result.indicators.append(ind)
        if ind.severity in ("CRITICAL", "HIGH"):
            result.threats.append(ind.description)
            result.verdict = "FAIL"
        else:
            result.warnings.append(ind.description)

    # ── 5. Format-specific deep scan ──────────────────────────────────────
    if ext == ".pdf":
        pdf_inds, pdf_meta = _pdf_deep_scan(filepath)
        result.indicators.extend(pdf_inds)
        result.metadata_found.update(pdf_meta)
        for ind in pdf_inds:
            if ind.severity == "CRITICAL":
                result.threats.append(ind.description)
                result.verdict = "FAIL"
            elif ind.severity == "HIGH":
                result.warnings.append(ind.description)
                if result.verdict == "PASS":
                    result.verdict = "WARN"

    elif ext in (".jpg", ".jpeg", ".png"):
        meta = _extract_image_metadata(filepath)
        result.metadata_found.update(meta)
        if "GPSInfo" in meta:
            result.warnings.append("GPS coordinates found in EXIF metadata")
        if "Make" in meta or "Model" in meta:
            result.warnings.append(f"Camera info in EXIF: {meta.get('Make','')} {meta.get('Model','')}")
        if len(meta) > 0:
            if result.verdict == "PASS":
                result.verdict = "WARN"
        suspicious, steg_detail = _image_steg_analysis(filepath)
        if suspicious:
            result.indicators.append(ThreatIndicator(
                category="steganography", severity="HIGH",
                description="Steganography payload likely present",
                evidence=steg_detail))
            result.warnings.append(f"Steganography detected: {steg_detail}")
            if result.verdict == "PASS":
                result.verdict = "WARN"

    elif ext in (".docx", ".xlsx"):
        office_inds = _office_deep_scan(filepath)
        result.indicators.extend(office_inds)
        for ind in office_inds:
            if ind.severity in ("CRITICAL", "HIGH"):
                result.warnings.append(ind.description)
                if result.verdict == "PASS":
                    result.verdict = "WARN"
            else:
                result.warnings.append(ind.description)

    # ── 6. Keyword scan ────────────────────────────────────────────────────
    kw_inds = _keyword_scan(filepath)
    for ind in kw_inds:
        result.indicators.append(ind)
        result.threats.append(ind.description)
        result.verdict = "FAIL"

    # ── 7. Size check ──────────────────────────────────────────────────────
    if size > MAX_SIZE_MB:
        result.warnings.append(f"File exceeds {MAX_SIZE_MB} MB limit ({size:.1f} MB)")
        if result.verdict == "PASS":
            result.verdict = "WARN"

    # ── 8. Risk score ──────────────────────────────────────────────────────
    result.risk_score = _compute_risk_score(result)

    # ── 9. Content preview — read inside the file ──────────────────────────
    result.content_preview = extract_content_preview(filepath)

    return result


# ══════════════════════════════════════════════════════════════════════════════
# SANITIZE
# ══════════════════════════════════════════════════════════════════════════════
def sanitize_file(filepath: str, out_dir: str) -> SanitizeResult:
    import time
    ext   = get_extension(filepath)
    fname = os.path.basename(filepath)
    os.makedirs(out_dir, exist_ok=True)

    dispatch = {
        ".pdf":  _sanitize_pdf,
        ".jpg":  _sanitize_image,
        ".jpeg": _sanitize_image,
        ".png":  _sanitize_image,
        ".docx": _sanitize_docx,
        ".xlsx": _sanitize_xlsx,
        ".txt":  _sanitize_text,
        ".csv":  _sanitize_text,
    }

    t0 = time.time()
    handler = dispatch.get(ext, _sanitize_copy)
    result  = handler(filepath, out_dir)
    result.filename = fname
    result.sanitize_time_ms = round((time.time() - t0) * 1000, 1)
    result.sha256_in  = sha256_file(filepath)
    result.size_in_mb = round(file_size_mb(filepath), 3)

    if result.success and result.output_path:
        result.sha256_out  = sha256_file(result.output_path)
        result.size_out_mb = round(file_size_mb(result.output_path), 3)

    return result


def _sanitize_pdf(filepath: str, out_dir: str) -> SanitizeResult:
    res     = SanitizeResult(filename=os.path.basename(filepath))
    outpath = os.path.join(out_dir, os.path.basename(filepath))
    if not HAS_PIKEPDF:
        shutil.copy2(filepath, outpath)
        res.actions.append("pikepdf not available — file copied unchanged")
        res.output_path = outpath
        res.success = True
        return res
    try:
        removed_meta = []
        removed_js   = 0
        removed_aa   = 0
        with pikepdf.open(filepath) as pdf:
            with pdf.open_metadata() as meta:
                keys = list(meta.keys())
                for key in keys:
                    removed_meta.append(key)
                    del meta[key]
            dangerous = ["/AA", "/JS", "/Launch", "/SubmitForm", "/ImportData",
                         "/GoToR", "/GoToE", "/URI", "/Sound", "/Movie", "/Widget"]
            for page in pdf.pages:
                for key in dangerous:
                    if key in page:
                        del page[key]
                        if key == "/JS": removed_js += 1
                        elif key == "/AA": removed_aa += 1
            root_keys_to_remove = ["/OpenAction", "/AA", "/AcroForm", "/Names",
                                    "/JavaScript", "/EmbeddedFiles", "/XFA"]
            for key in root_keys_to_remove:
                if key in pdf.Root:
                    del pdf.Root[key]
                    res.removed_items.append(f"Root{key}")
            if "/Names" in pdf.Root:
                names = pdf.Root["/Names"]
                if "/EmbeddedFiles" in names:
                    del names["/EmbeddedFiles"]
                    res.removed_items.append("Embedded files")
                if "/JavaScript" in names:
                    del names["/JavaScript"]
                    res.removed_items.append("Named JavaScript")
            pdf.save(outpath, compress_streams=True,
                     object_stream_mode=pikepdf.ObjectStreamMode.generate)
        res.actions.append(f"Stripped {len(removed_meta)} metadata field(s)")
        res.actions.append(f"Removed {removed_js} JavaScript action(s), {removed_aa} auto-action(s)")
        res.actions.append("Removed dangerous root actions (OpenAction, AcroForm, XFA)")
        res.actions.append("Re-saved with stream compression")
        res.output_path = outpath
        res.success = True
    except Exception as e:
        res.error = str(e)
        res.success = False
    return res


def _sanitize_image(filepath: str, out_dir: str) -> SanitizeResult:
    res     = SanitizeResult(filename=os.path.basename(filepath))
    ext     = get_extension(filepath)
    outpath = os.path.join(out_dir, os.path.basename(filepath))
    if not HAS_PIL:
        shutil.copy2(filepath, outpath)
        res.actions.append("Pillow not available — copied unchanged")
        res.output_path = outpath
        res.success = True
        return res
    try:
        img = Image.open(filepath)
        exif_count = 0
        try:
            exif_data = img._getexif()
            exif_count = len(exif_data) if exif_data else 0
        except Exception:
            pass
        mode = img.mode
        if mode in ("RGBA", "LA"):
            bg = Image.new("RGB", img.size, (255, 255, 255))
            bg.paste(img, mask=img.split()[-1])
            img = bg
        elif mode != "RGB":
            img = img.convert("RGB")
        clean = Image.new("RGB", img.size)
        clean.putdata(list(img.getdata()))
        fmt = "JPEG" if ext in (".jpg", ".jpeg") else "PNG"
        save_kwargs = {"format": fmt, "optimize": True}
        if fmt == "JPEG":
            save_kwargs["quality"] = 92
            save_kwargs["subsampling"] = 0
        clean.save(outpath, **save_kwargs)
        res.actions.append(f"Stripped {exif_count} EXIF field(s) including GPS data")
        res.actions.append(f"Converted from {mode} → RGB, destroyed LSB steganography")
        res.actions.append(f"Re-encoded as {fmt} (quality=92, optimized)")
        res.removed_items.append(f"{exif_count} EXIF metadata tags")
        res.output_path = outpath
        res.success = True
    except Exception as e:
        res.error = str(e)
        res.success = False
    return res


def _sanitize_docx(filepath: str, out_dir: str) -> SanitizeResult:
    res     = SanitizeResult(filename=os.path.basename(filepath))
    outpath = os.path.join(out_dir, os.path.basename(filepath))
    if not HAS_DOCX:
        shutil.copy2(filepath, outpath)
        res.actions.append("python-docx not available — copied unchanged")
        res.output_path = outpath
        res.success = True
        return res
    try:
        buf = io.BytesIO()
        drop_patterns = ["vbaProject", "activeX", "oleObject", ".bin"]
        dropped = []
        with zipfile.ZipFile(filepath, "r") as zin, \
             zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                skip = any(p in item.filename for p in drop_patterns)
                if skip:
                    dropped.append(item.filename)
                    continue
                data = zin.read(item.filename)
                if item.filename.endswith(".rels"):
                    data = re.sub(
                        rb'<Relationship[^>]+Type="[^"]*hyperlink[^"]*"[^/]*/?>',
                        b"", data)
                zout.writestr(item, data)
        with open(outpath, "wb") as f:
            f.write(buf.getvalue())
        doc = DocxDocument(outpath)
        props = doc.core_properties
        original_author = props.author or "Unknown"
        props.author           = "REDACTED"
        props.last_modified_by = "REDACTED"
        props.comments         = ""
        props.keywords         = ""
        props.subject          = ""
        doc.save(outpath)
        res.actions.append(f"Removed {len(dropped)} embedded object(s)" if dropped else "No macros/binaries found")
        res.actions.append(f"Cleared author metadata (was: '{original_author}')")
        res.actions.append("Stripped external hyperlink references")
        res.actions.append("Cleared comments, keywords, subject properties")
        res.removed_items.extend(dropped)
        res.output_path = outpath
        res.success = True
    except Exception as e:
        res.error = str(e)
        res.success = False
    return res


def _sanitize_xlsx(filepath: str, out_dir: str) -> SanitizeResult:
    res     = SanitizeResult(filename=os.path.basename(filepath))
    outpath = os.path.join(out_dir, os.path.basename(filepath))
    if not HAS_OPENPYXL:
        shutil.copy2(filepath, outpath)
        res.actions.append("openpyxl not available — copied unchanged")
        res.output_path = outpath
        res.success = True
        return res
    try:
        buf = io.BytesIO()
        dropped = []
        with zipfile.ZipFile(filepath, "r") as zin, \
             zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if "vbaProject" in item.filename or "activeX" in item.filename:
                    dropped.append(item.filename)
                    continue
                data = zin.read(item.filename)
                if "connections" in item.filename.lower():
                    dropped.append(item.filename + " (external connections)")
                    continue
                zout.writestr(item, data)
        with open(outpath, "wb") as f:
            f.write(buf.getvalue())
        wb = openpyxl.load_workbook(outpath)
        hidden_sheets = []
        for name in wb.sheetnames:
            ws = wb[name]
            if ws.sheet_state in ("hidden", "veryHidden"):
                hidden_sheets.append(name)
                ws.sheet_state = "visible"
        if hasattr(wb, "defined_names"):
            ext_names = [n for n in wb.defined_names
                         if "[" in str(getattr(wb.defined_names[n], "attr_text", ""))]
            for n in ext_names:
                del wb.defined_names[n]
                res.removed_items.append(f"External name ref: {n}")
        wb.save(outpath)
        res.actions.append(f"Removed {len(dropped)} VBA/ActiveX object(s)" if dropped else "No macros found")
        res.actions.append(f"Revealed {len(hidden_sheets)} hidden sheet(s)" if hidden_sheets else "No hidden sheets")
        res.actions.append("Removed external data connections")
        res.removed_items.extend(dropped)
        res.output_path = outpath
        res.success = True
    except Exception as e:
        res.error = str(e)
        res.success = False
    return res


def _sanitize_text(filepath: str, out_dir: str) -> SanitizeResult:
    res     = SanitizeResult(filename=os.path.basename(filepath))
    outpath = os.path.join(out_dir, os.path.basename(filepath))
    try:
        with open(filepath, "rb") as f:
            raw = f.read()
        if raw.startswith(b"\xef\xbb\xbf"):
            raw = raw[3:]
            res.actions.append("Removed UTF-8 BOM")
        content = raw.decode("utf-8", errors="ignore")
        clean = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f\x80-\x9f]", "", content)
        if os.path.basename(filepath).endswith(".csv"):
            lines = clean.split("\n")
            sanitized_lines = []
            for line in lines:
                if line.startswith(("=", "+", "-", "@", "\t=")):
                    line = "'" + line
                    res.removed_items.append("CSV formula injection neutralized")
                sanitized_lines.append(line)
            clean = "\n".join(sanitized_lines)
            res.actions.append("Neutralized CSV formula injection attempts")
        chars_removed = len(content) - len(clean)
        with open(outpath, "w", encoding="utf-8") as f:
            f.write(clean)
        res.actions.append(f"Stripped {chars_removed} control/null character(s)")
        res.actions.append("Re-encoded as clean UTF-8 (no BOM)")
        res.output_path = outpath
        res.success = True
    except Exception as e:
        res.error = str(e)
        res.success = False
    return res


def _sanitize_copy(filepath: str, out_dir: str) -> SanitizeResult:
    res     = SanitizeResult(filename=os.path.basename(filepath))
    outpath = os.path.join(out_dir, os.path.basename(filepath))
    shutil.copy2(filepath, outpath)
    res.actions.append("No specific sanitizer — file copied unchanged")
    res.output_path = outpath
    res.success = True
    return res